#!/usr/bin/env python3
"""
Extract evidence images and text from patent claim chart source documents.

Handles DOCX, XLSX, and PPTX files. Outputs a structured JSON report and
copies unique images to a staging directory for use in claim chart generation.

Usage:
    python3 extract_evidence.py <project_dir> <output_dir> [--patent PATENT_ID]

Example:
    python3 extract_evidence.py /path/to/project /tmp/evidence --patent US00000000
"""

import argparse
import hashlib
import json
import os
import shutil
import struct
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


def get_image_info(data):
    """Get image format and dimensions from raw bytes."""
    if data[:4] == b'\x89PNG' and len(data) >= 24:
        w = struct.unpack('>I', data[16:20])[0]
        h = struct.unpack('>I', data[20:24])[0]
        return {"format": "PNG", "width": w, "height": h}
    elif data[:2] == b'\xff\xd8':
        return {"format": "JPEG", "width": 0, "height": 0}
    return {"format": "UNKNOWN", "width": 0, "height": 0}


def extract_media_from_zip(filepath, output_dir):
    """Extract media files from a DOCX/PPTX zip archive."""
    images = []
    tag = Path(filepath).stem.replace(" ", "_")
    dest = os.path.join(output_dir, tag)
    os.makedirs(dest, exist_ok=True)

    try:
        with zipfile.ZipFile(filepath) as z:
            media_files = [
                f for f in z.namelist()
                if "media/" in f and any(f.lower().endswith(ext) for ext in
                    ('.png', '.jpg', '.jpeg', '.tmp', '.gif', '.bmp'))
            ]
            for mf in media_files:
                z.extract(mf, dest)
                fpath = os.path.join(dest, mf)
                with open(fpath, 'rb') as fh:
                    data = fh.read()
                content_hash = hashlib.md5(data).hexdigest()
                info = get_image_info(data)
                images.append({
                    "path": fpath,
                    "filename": os.path.basename(mf),
                    "source_doc": os.path.basename(filepath),
                    "size_bytes": len(data),
                    "hash": content_hash,
                    **info
                })
    except Exception as e:
        print(f"  Warning: Could not extract from {filepath}: {e}", file=sys.stderr)

    return images


def extract_docx_evidence(filepath):
    """Extract text evidence and image-to-cell mapping from a DOCX claim chart."""
    if DocxDocument is None:
        return None

    try:
        doc = DocxDocument(filepath)
    except Exception as e:
        print(f"  Warning: Could not read {filepath}: {e}", file=sys.stderr)
        return None

    if not doc.tables:
        return None

    result = {
        "file": os.path.basename(filepath),
        "tables": len(doc.tables),
        "rows": []
    }

    # Usually the main claim chart is the last (or only) multi-row table
    main_table = None
    for table in doc.tables:
        if len(table.rows) > 2:
            main_table = table
            break
    if main_table is None:
        main_table = doc.tables[-1]

    num_cols = len(main_table.columns)

    for ri, row in enumerate(main_table.rows):
        row_data = {"row_index": ri, "cells": []}
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            # Count images in cell
            img_count = 0
            for para in cell.paragraphs:
                for run in para.runs:
                    blips = run._element.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    img_count += len(blips)

            # Extract URLs
            urls = []
            import re
            for url_match in re.finditer(r'https?://[^\s)]+', text):
                urls.append(url_match.group())

            col_name = "unknown"
            if num_cols == 3:
                col_name = ["claim", "spec_support", "evidence"][ci] if ci < 3 else f"col{ci}"
            elif num_cols == 2:
                col_name = ["claim", "evidence"][ci] if ci < 2 else f"col{ci}"

            row_data["cells"].append({
                "column": col_name,
                "col_index": ci,
                "text": text[:2000],
                "image_count": img_count,
                "urls": urls
            })
        result["rows"].append(row_data)

    return result


def extract_xlsx_mapping(filepath):
    """Extract mapping data from Excel files."""
    if openpyxl is None:
        return None

    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
    except Exception as e:
        print(f"  Warning: Could not read {filepath}: {e}", file=sys.stderr)
        return None

    result = {"file": os.path.basename(filepath), "sheets": []}

    for sname in wb.sheetnames:
        ws = wb[sname]
        sheet_data = {
            "name": sname,
            "rows": ws.max_row,
            "cols": ws.max_column,
            "data": []
        }
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 50), values_only=False):
            vals = []
            for c in row:
                v = str(c.value) if c.value is not None else ""
                vals.append(v[:500])
            if any(v for v in vals):
                sheet_data["data"].append(vals)
        result["sheets"].append(sheet_data)

    return result


def deduplicate_images(all_images):
    """Deduplicate images by content hash, keeping the first occurrence."""
    seen = {}
    unique = []
    for img in all_images:
        h = img["hash"]
        if h not in seen:
            seen[h] = img
            unique.append(img)
        else:
            seen[h].setdefault("also_in", []).append(img["source_doc"])
    return unique


def stage_images(unique_images, staging_dir):
    """Copy unique images to a flat staging directory with sequential names."""
    os.makedirs(staging_dir, exist_ok=True)
    staged = []
    for i, img in enumerate(unique_images):
        ext = ".png" if img["format"] == "PNG" else ".jpg" if img["format"] == "JPEG" else ".bin"
        dest_name = f"img_{i:03d}{ext}"
        dest_path = os.path.join(staging_dir, dest_name)
        shutil.copy2(img["path"], dest_path)
        img["staged_path"] = dest_path
        img["staged_name"] = dest_name
        staged.append(img)
    return staged


def main():
    parser = argparse.ArgumentParser(description="Extract evidence from claim chart project")
    parser.add_argument("project_dir", help="Project directory containing source documents")
    parser.add_argument("output_dir", help="Output directory for extracted evidence")
    parser.add_argument("--patent", help="Filter for specific patent ID (e.g., US00000000)")
    args = parser.parse_args()

    project = Path(args.project_dir)
    output = Path(args.output_dir)
    output.mkdir(parents=True, exist_ok=True)

    # Collect all relevant files
    docx_files = []
    xlsx_files = []
    pptx_files = []

    for root, dirs, files in os.walk(project):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for f in files:
            fpath = os.path.join(root, f)
            fl = f.lower()
            if fl.endswith('.docx') and not fl.startswith('~$'):
                docx_files.append(fpath)
            elif fl.endswith('.xlsx') and not fl.startswith('~$'):
                xlsx_files.append(fpath)
            elif fl.endswith('.pptx') and not fl.startswith('~$'):
                pptx_files.append(fpath)

    print(f"Found: {len(docx_files)} DOCX, {len(xlsx_files)} XLSX, {len(pptx_files)} PPTX")

    # Extract images from all documents
    all_images = []
    for f in docx_files + pptx_files:
        imgs = extract_media_from_zip(f, str(output / "raw_media"))
        if imgs:
            print(f"  {os.path.basename(f)}: {len(imgs)} images")
            all_images.extend(imgs)

    # Deduplicate
    unique = deduplicate_images(all_images)
    print(f"\nTotal images: {len(all_images)}, Unique: {len(unique)}")

    # Filter to PNG only (for docx generation compatibility)
    png_images = [img for img in unique if img["format"] == "PNG"]
    print(f"PNG images: {len(png_images)}")

    # Sort by size (larger images tend to be more informative)
    png_images.sort(key=lambda x: x["size_bytes"], reverse=True)

    # Stage images
    staging = str(output / "staged")
    staged = stage_images(png_images, staging)

    # Extract text evidence from DOCX files
    evidence_data = []
    for f in docx_files:
        ev = extract_docx_evidence(f)
        if ev:
            evidence_data.append(ev)

    # Extract Excel mapping data
    mapping_data = []
    for f in xlsx_files:
        mp = extract_xlsx_mapping(f)
        if mp:
            mapping_data.append(mp)

    # Collect all unique URLs
    all_urls = set()
    for ev in evidence_data:
        for row in ev["rows"]:
            for cell in row["cells"]:
                for url in cell.get("urls", []):
                    all_urls.add(url)

    # Write report
    report = {
        "project_dir": str(project),
        "files_scanned": {
            "docx": len(docx_files),
            "xlsx": len(xlsx_files),
            "pptx": len(pptx_files)
        },
        "images": {
            "total_extracted": len(all_images),
            "unique": len(unique),
            "png_only": len(png_images),
            "staged_dir": staging
        },
        "staged_images": [
            {
                "staged_path": img["staged_path"],
                "staged_name": img["staged_name"],
                "source_doc": img["source_doc"],
                "dimensions": f"{img['width']}x{img['height']}",
                "size_kb": img["size_bytes"] // 1024,
                "also_in": img.get("also_in", [])
            }
            for img in staged
        ],
        "evidence_documents": evidence_data,
        "mapping_data": mapping_data,
        "all_urls": sorted(all_urls),
        "url_count": len(all_urls)
    }

    report_path = str(output / "evidence_report.json")
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    print(f"\nReport saved: {report_path}")
    print(f"Staged images: {staging}/")
    print(f"Unique URLs: {len(all_urls)}")

    return report


if __name__ == "__main__":
    main()
